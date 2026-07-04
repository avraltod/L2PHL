"""Fill L2PHL weighting-note tables from the diagnostics workbook.

Caption-anchored and prose-safe: for each managed (caption, sheet) pair it finds
the Heading-2 caption paragraph, takes the table immediately following it, and
rewrites that table's cells from weight_diagnostics_results.xlsx. Prose is never
touched. Every number comes from the workbook - none is hand-typed. Applies the
KGZ house table style (Arial Narrow 9pt, bold header, Table Grid, tight rows).
Re-run after the panel rolls forward and the diagnostics do-file regenerates the
workbook.
"""
import openpyxl
from docx import Document
from docx.table import Table
from docx.shared import Pt
from docx.oxml.ns import qn

DOCX = "CAPI/Analysis/SL/doc/L2PHL_WEIGHTING_TECHNICAL_NOTE.docx"
XLSX = "CATI/Analysis/do/output/weight_diagnostics_results.xlsx"

# (Heading-2 caption in the note, results sheet name) in document order
MANAGED = [
    ("CAPI baseline: weight distribution and design effect", "capi_dist"),
    ("CAPI baseline: effect of calibration (design vs final weights)", "capi_caleffect"),
    ("CATI panel: individual weight by round", "cati_indw"),
    ("CATI panel: household and population weight by round", "cati_hhw"),
    ("CATI panel: composition stability by round", "cati_attrition"),
    ("Retention by round", "attr_retention"),
    ("Attrition: stayers versus leavers", "attr_bias"),
]

FONT, SIZE = "Arial Narrow", Pt(9)

# presentation rules
DROP_ROWS = {"attr_bias": {"employed"}}                 # R1 employment too thin to test
RELABEL   = {"attr_bias": {"region(via stratum)": "geography (stratum)"}}
P_COLS    = {"p_value", "p"}                             # rendered to 3 decimals


def sheet_rows(ws):
    return [[c.value for c in r] for r in ws.iter_rows()
            if any(c.value is not None for c in r)]


def fmt(header, value):
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    try:
        f = float(value)
    except (TypeError, ValueError):
        return str(value)
    if header in P_COLS:
        return f"{f:.3f}"
    if f == int(f):
        return f"{int(f):,}"
    if abs(f) >= 1000:
        return f"{f:,.0f}"
    if abs(f) >= 100:
        return f"{f:,.1f}"
    return f"{f:.2f}"


def style_table(tbl, doc):
    tbl.style = doc.styles["Table Grid"]
    for ri, row in enumerate(tbl.rows):
        for c in row.cells:
            for p in c.paragraphs:
                pf = p.paragraph_format
                pf.space_before = Pt(1); pf.space_after = Pt(1); pf.line_spacing = 1.0
                if not p.runs and p.text:
                    p.add_run("")
                for r in p.runs:
                    r.font.name = FONT; r.font.size = SIZE
                    if ri == 0:
                        r.font.bold = True


def next_table_el(caption_p):
    """First <w:tbl> after the caption paragraph, stopping at the next heading."""
    el = caption_p._p.getnext()
    while el is not None:
        if el.tag == qn("w:tbl"):
            return el
        if el.tag == qn("w:p"):
            pPr = el.find(qn("w:pPr"))
            if pPr is not None:
                ps = pPr.find(qn("w:pStyle"))
                if ps is not None and str(ps.get(qn("w:val"))).startswith("Heading"):
                    return None
        el = el.getnext()
    return None


def main():
    wb = openpyxl.load_workbook(XLSX, data_only=True)
    doc = Document(DOCX)
    caps = {p.text.strip(): p for p in doc.paragraphs
            if p.style.name.startswith("Heading 2")}

    filled = 0
    for caption, sheet in MANAGED:
        cp = caps.get(caption)
        if cp is None:
            print(f"  [skip] caption not found: {caption!r}")
            continue
        rows = sheet_rows(wb[sheet])
        hdr, body = rows[0], rows[1:]
        drop = DROP_ROWS.get(sheet, set())
        relab = RELABEL.get(sheet, {})
        body = [r for r in body if str(r[0]) not in drop]
        for r in body:
            if str(r[0]) in relab:
                r[0] = relab[str(r[0])]
        data = [hdr] + body

        tel = next_table_el(cp)
        if tel is None:
            tbl = doc.add_table(rows=len(data), cols=len(hdr))
            cp._p.addnext(tbl._element)
        else:
            tbl = Table(tel, cp._parent)
            while len(tbl.rows) > 1:
                tbl.rows[-1]._element.getparent().remove(tbl.rows[-1]._element)
            while len(tbl.rows) < len(data):
                tbl.add_row()

        for i, rw in enumerate(data):
            for j, val in enumerate(rw):
                if j < len(tbl.rows[i].cells):
                    tbl.rows[i].cells[j].text = str(val) if i == 0 else fmt(hdr[j], val)
        style_table(tbl, doc)
        filled += 1

    doc.save(DOCX)
    print(f"filled {filled}/{len(MANAGED)} tables from {XLSX}")


if __name__ == "__main__":
    main()
